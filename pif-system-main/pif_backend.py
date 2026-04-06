import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import docx
import re
import os
import io

# ─────────────────────────────────────────────
# UTILITY: Sanitize filename (Fix #7)
# ─────────────────────────────────────────────
def sanitize_filename(name: str) -> str:
    """Loại bỏ ký tự không hợp lệ trong tên file Windows."""
    name = re.sub(r'[\\/:*?"<>|]', '_', str(name))
    name = re.sub(r'\s+', ' ', name).strip()
    return name[:80]  # giới hạn độ dài tên file


def format_date_vietnamese(date_str):
    """Convert DD/MM/YYYY to 'ngày DD tháng MM năm YYYY'"""
    try:
        if isinstance(date_str, pd.Timestamp):
            day, month, year = date_str.day, date_str.month, date_str.year
        else:
            match = re.search(r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})', str(date_str))
            if match:
                day, month, year = match.groups()
            else:
                return str(date_str)
        return f"ngày {day} tháng {month} năm {year}"
    except Exception:
        return str(date_str)


def extract_product_info(df):
    """
    Extract product metadata from the formula Excel (ĐMVT).
    Robust for different Excel structures.
    """
    info = {
        'ten_san_pham': '',
        'ngay_ban_hanh': '',
        'dang_bao_che': '',
        'quy_cach_dong_goi': '',
        'quy_cach': '',
        'ma_so_pif': '',
    }

    # 1. Ngày ban hành
    for r in range(min(10, len(df))):
        for c in range(min(10, len(df.columns))):
            val = str(df.iloc[r, c])
            if "Ngày ban hành:" in val:
                date_match = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{4})', val)
                if date_match:
                    info['ngay_ban_hanh'] = format_date_vietnamese(date_match.group(1))
                break

    # 2. Tên sản phẩm
    for r in range(min(10, len(df))):
        val = str(df.iloc[r, 0])
        if "Tên sản phẩm:" in val:
            info['ten_san_pham'] = val.split("Tên sản phẩm:")[1].strip()
            break

    # 3. Quy cách đóng gói
    for r in range(min(10, len(df))):
        for c in range(min(5, len(df.columns))):
            val = str(df.iloc[r, c])
            if "Quy cách" in val:
                spec_match = re.search(r'Quy cách\s*:?\s*(.*)', val, re.IGNORECASE)
                if spec_match:
                    spec = spec_match.group(1).strip()
                    info['quy_cach_dong_goi'] = spec
                    info['quy_cach'] = spec
                    break
        if info['quy_cach_dong_goi']:
            break

    # 4. Dạng bào chế
    for r in range(min(10, len(df))):
        val = str(df.iloc[r, 0])
        if "Dạng bào chế:" in val:
            info['dang_bao_che'] = val.split("Dạng bào chế:")[1].strip()
            break

    # 5. Mã số PIF (ĐMVT Code)
    # Tìm kiếm rộng hơn (10 dòng đầu, tất cả cột) để bao quát ô gộp
    for r in range(min(10, len(df))):
        for c in range(len(df.columns)):
            val = str(df.iloc[r, c])
            if "Mã số:" in val or "Mã sản phẩm:" in val:
                # Regex tìm chuỗi số dài (ít nhất 4 chữ số), có thể có gạch ngang/chấm
                id_match = re.search(r'(\d{4,})', val)
                if id_match:
                    full_id = id_match.group(1)
                    # Lấy 3 số cuối của mã vật tư
                    info['ma_so_pif'] = full_id[-3:]
                break
        if info['ma_so_pif']:
            break

    return info


def extract_ingredients(df):
    """
    Extract ingredient table from formula Excel (ĐMVT).
    """
    ingredients = []
    try:
        for i in range(10, len(df)):
            row = df.iloc[i]
            stt = row[0]
            if pd.isna(stt) or str(stt).strip() == "":
                continue
            if "Tổng" in str(stt) or "Bao bì" in str(stt):
                if "Bao bì" not in str(stt):
                    break
                continue
            try:
                val_stt = float(stt)
                if val_stt.is_integer():
                    inci_name = str(row[3]).strip() if not pd.isna(row[3]) else ""
                    if inci_name == "":
                        continue
                    ingredients.append({
                        'ma_vt':   str(row[1]).strip() if not pd.isna(row[1]) else "",
                        'ten_vt':  str(row[2]).strip() if not pd.isna(row[2]) else "",
                        'ten_inci': inci_name,
                        'ti_le':   "",
                        'ghi_chu': ""
                    })
            except ValueError:
                continue
    except Exception as e:
        print(f"Error extracting ingredients: {e}")
    return ingredients


def _find_db_key_column(db_df: pd.DataFrame) -> str:
    """
    Tìm cột mã nguyên liệu trong DB một cách an toàn, tránh crash khi
    tên cột bị lỗi encoding (Fix #12 – tên cột 'ương').
    Ưu tiên: 'Mã vật tư', 'Ma vat tu', 'ương', hoặc cột đầu tiên.
    """
    candidates = ['Mã vật tư', 'Ma vat tu', 'ma_vt', 'MÃ VẬT TƯ', 'ương']
    for name in candidates:
        if name in db_df.columns:
            return name
    # Fallback: cột đầu tiên
    return db_df.columns[0]


def merge_ingredient_data(ingredients, db_df):
    """
    Merge ingredients với dữ liệu an toàn & công dụng từ DB T01.
    """
    key_col = _find_db_key_column(db_df)

    # Tìm cột công dụng/an toàn linh hoạt
    def find_col(df, keywords):
        for col in df.columns:
            col_lower = str(col).lower()
            if any(k.lower() in col_lower for k in keywords):
                return col
        return None

    col_short   = find_col(db_df, ['ngắn gọn', 'ngan gon', 'short'])
    col_safety  = find_col(db_df, ['an toàn', 'an toan', 'safety', 'dữ liệu an'])
    col_detail  = find_col(db_df, ['chi tiết', 'chi tiet', 'detail'])

    db_dict = {}
    for _, row in db_df.iterrows():
        key = str(row[key_col]).strip()
        db_dict[key] = {
            'short_usage':  str(row[col_short]).strip()  if col_short  and not pd.isna(row[col_short])  else "",
            'safety_data':  str(row[col_safety]).strip() if col_safety and not pd.isna(row[col_safety]) else "",
            'detail_usage': str(row[col_detail]).strip() if col_detail and not pd.isna(row[col_detail]) else "",
        }

    merged = []
    for ing in ingredients:
        ma = ing['ma_vt']
        db_info = db_dict.get(ma, {'short_usage': '', 'safety_data': '', 'detail_usage': ''})
        if ma == "" and "Water" in ing['ten_inci']:
            db_info = {
                'short_usage':  'Dung môi',
                'safety_data':  'An toàn',
                'detail_usage': 'Dung môi hòa tan các hoạt chất'
            }
        merged.append({**ing, **db_info})
    return merged


def set_text_style(element, font_name="Times New Roman", size=13):
    """Utility: set font và size cho paragraph/cell."""
    paragraphs = element.paragraphs if hasattr(element, 'paragraphs') else [element]
    for paragraph in paragraphs:
        if not paragraph.runs:
            paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)


def update_210_safety_details(doc, ingredients):
    """Logic cho 2.10: Thay thế placeholder bằng dữ liệu an toàn thực."""
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Dữ liệu an toàn của nguyên liệu" in p.text:
            start_idx = i
            break
    if start_idx == -1:
        return

    paras_to_remove = doc.paragraphs[start_idx + 1:]
    parent = doc.paragraphs[start_idx]._element.getparent()
    for p in paras_to_remove:
        parent.remove(p._element)

    last_para = doc.paragraphs[start_idx]
    current_ref = last_para

    for i, ing in enumerate(ingredients, 1):
        # Header
        p_head = docx.oxml.OxmlElement('w:p')
        current_ref._element.addnext(p_head)
        current_ref = docx.text.paragraph.Paragraph(p_head, last_para._parent)
        run_head = current_ref.add_run(f"{i}. {ing['ten_inci']}")
        run_head.bold = True
        set_text_style(current_ref)

        # Body
        p_body = docx.oxml.OxmlElement('w:p')
        current_ref._element.addnext(p_body)
        current_ref = docx.text.paragraph.Paragraph(p_body, last_para._parent)
        safety_txt = ing['safety_data'] if ing['safety_data'] else "Chưa có dữ liệu"
        current_ref.add_run(safety_txt)
        set_text_style(current_ref)


def fill_word_table(doc, ingredients, mode='short'):
    """
    Tìm bảng đúng và điền dữ liệu nguyên liệu.
    Mapping cột 5 cột: STT | Tên INCI | Tỉ lệ | Ghi chú | Công dụng
    """
    target_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            header_found = False
            for r_idx in range(min(2, len(table.rows))):
                row_text = "".join([c.text.lower() for c in table.rows[r_idx].cells])
                if "stt" in row_text:
                    header_found = True
                    break
            if header_found:
                target_table = table
                break

    if not target_table:
        if len(doc.tables) > 1:
            target_table = doc.tables[1]
        elif len(doc.tables) > 0:
            target_table = doc.tables[0]
        else:
            return

    table = target_table

    header_row_count = 1
    for r_idx in range(min(3, len(table.rows))):
        row_text = "".join([c.text.lower() for c in table.rows[r_idx].cells])
        if "stt" in row_text:
            header_row_count = r_idx + 1

    current_row_idx = header_row_count

    for i, ing in enumerate(ingredients, 1):
        if current_row_idx < len(table.rows):
            row = table.rows[current_row_idx]
            if "".join([c.text.strip() for c in row.cells]) == "":
                row_cells = row.cells
            else:
                row_cells = table.add_row().cells
        else:
            row_cells = table.add_row().cells

        current_row_idx += 1

        row_cells[0].text = str(i)
        row_cells[1].text = ing['ten_inci']

        if len(row_cells) > 2:
            row_cells[2].text = ing['ti_le']
        if len(row_cells) > 3:
            row_cells[3].text = ing['ghi_chu']

        usage_val = ""
        if mode in ('short', 'safety'):
            usage_val = ing['short_usage']
        elif mode == 'detail':
            usage_val = ing['detail_usage']

        if len(row_cells) > 4:
            row_cells[4].text = usage_val
        elif len(row_cells) == 3:
            row_cells[2].text = usage_val

        for cell in row_cells:
            set_text_style(cell)

    # Xóa các dòng thừa
    while len(table.rows) > current_row_idx:
        last_row = table.rows[-1]
        table._element.remove(last_row._element)


def process_single_formula(formula_file, db_df, templates_dir):
    """
    Xử lý 1 file ĐMVT và tạo 7 tài liệu PIF.
    Trả về dict {tên_file: bytes}.
    """
    df = pd.read_excel(formula_file, header=None)
    metadata = extract_product_info(df)
    ing_list = extract_ingredients(df)
    merged_ingredients = merge_ingredient_data(ing_list, db_df)

    results = {}

    originals = [
        '1.4_ban_cong_bo_phu_hop_quy_dinh_cgmp.docx',
        '1.5_cong_bo_an_toan_suc_khoe.docx',
        '1.6_tom_tat_tac_dung_khong_mong_muon.docx',
        '1.7_bao_cao_tinh_nang_cong_dung.docx',
        '2.10_du_lieu_an_toan_cua_nguyen_lieu.docx',
        '4.15_bao_cao_danh_gia_tinh_an_toan.docx',
        '4.18.1_bao_cao_day_du_tinh_nang_cong_dung.docx'
    ]

    all_files = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
    template_files = [f for f in all_files if f in originals]

    for t_name in template_files:
        t_path = os.path.join(templates_dir, t_name)

        mode = None
        if any(x in t_name for x in ['1.6', '1.7']):
            mode = 'short'
        elif any(x in t_name for x in ['4.15', '4.18.1']):
            mode = 'detail'
        elif '2.10' in t_name:
            mode = 'safety'

        tpl = DocxTemplate(t_path)

        context = metadata.copy()
        if any(x in t_name for x in ['4.15', '4.18.1']):
            context.pop('ngay_ban_hanh', None)
        if '2.10' in t_name:
            if 'ngay_ban_hanh' in context:
                context['ngay_ban_hanh'] = str(context['ngay_ban_hanh']).replace('ngày', 'Ngày', 1)

        tpl.render(context)

        if mode:
            fill_word_table(tpl.docx, merged_ingredients, mode=mode)
        if '2.10' in t_name:
            update_210_safety_details(tpl.docx, merged_ingredients)

        out_io = io.BytesIO()
        tpl.save(out_io)
        out_io.seek(0)

        prefix = t_name.split('_')[0]
        ma_so = metadata.get('ma_so_pif', '')
        product_name_safe = sanitize_filename(metadata['ten_san_pham'])

        # Định dạng file: [Prefix] [Mã 3 số] [Tên SP]
        # Nếu không có mã, chỉ dùng [Prefix] [Tên SP]
        if ma_so:
            fname_base = f"{prefix} {ma_so} {product_name_safe}"
        else:
            fname_base = f"{prefix} {product_name_safe}"

        out_filename = sanitize_filename(fname_base) + ".docx"
        results[out_filename] = out_io.read()

    return results, metadata
